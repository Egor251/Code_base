import docx
# https://docs-python.ru/packages/modul-python-docx-python/klass-document/


class Word:
    doc = docx.Document()
    head = doc.add_heading('Название файла')

    # Создаём таблицу
    def create_table(self, rows, cols):
        table = self.doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'
        table.autofit = True
        return table

    # Заполняем таблицу данными
    def fill_table(self, table, data):
        for i in range(len(data)):
            for j in range(len(data[0])):
                cell = table.cell(i, j)
                cell.text = str(data[i][j])

    # Закрываем документ и сохраняем его в указанный путь
    def finish(self, path):
        self.doc.save(path)
