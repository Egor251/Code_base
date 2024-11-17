"""
Пример заполнения шаблона docx при помощи библиотеки docxtpl.
"""

from docxtpl import DocxTemplate, InlineImage
from docx.shared import Mm

import inserted_data_example as insert_data  # Вставляемые данные в шаблон

import pandas as pd
import jinja2


class DocCreator:
    """
    Класс для создания документа Word с использованием шаблона и переданных данных.
    """

    counts = None
    lists = None  # bullet_list = [{'item': 'Мониторинг','sub_items': ['SSH – 16','RDP – 1','SNMP – 28']}
    tables = None
    images = None

    # Шаблон
    doc = DocxTemplate('template.docx')

    def __init__(self):
        self.counts = insert_data.counts
        self.lists = insert_data.lists
        self.tables = insert_data.tables
        self.images = insert_data.images

    def create_table(self, df: pd.DataFrame, heading: str = ''):
        """
        Создание таблицы в документе.

        :param df: DataFrame с данными
        :param heading: заголовок таблицы

        :return: None
        """

        self.doc.add_heading(heading, level=1)

        # Добавление таблицы в документ
        table = self.doc.add_table(rows=1, cols=len(df.columns))

        # Добавление заголовков
        hdr_cells = table.rows[0].cells
        for i, column_name in enumerate(df.columns):
            hdr_cells[i].text = column_name

        # Заполнение таблицы данными из DataFrame
        for index, row in df.iterrows():
            row_cells = table.add_row().cells
            for i, value in enumerate(row):
                row_cells[i].text = str(value)

    def read_docx(self, file_path: str) -> None:
        # Загружаем шаблон
        self.doc = DocxTemplate(file_path)

    @staticmethod
    def get_table_name(table):
        first_row_text = [cell.text for cell in table.rows[0].cells]
        print(f"Название таблицы {first_row_text[0]}")
        return first_row_text[0]

    def generate_docx(self, output_path):
        """
        Генерация документа с использованием шаблона.

        :param output_path: путь для сохранения готового документа

        :return: None
        :raises FileNotFoundError: если шаблон не найден
        :raises PermissionError: если нет прав на запись в указанный путь
        :raises Exception: если при создании документа произошла ошибка
        :raises ValueError: если в контексте не найдены данные для замены
        :raises TypeError: если в контексте не найдены данные для замены
        :raises KeyError: если в контексте не найдены данные для замены
        :raises AttributeError: если в контексте не найдены данные для
        """

        # Данные для замены
        context = {}

        # ИЗОБРАЖЕНИЯ! https://qna.habr.com/q/1195152
        for image in self.images.keys():
            # context[image] = InlineImage(self.doc,f'images/{self.images[image]}', width=Mm(100))
            # https://python-docx-template.readthedocs.io/en/latest/inline-images.html

            context[image] = InlineImage(self.doc, self.images[image], width=Mm(150))

        # Добавляем данные из переменных
        context.update(self.counts)
        context.update(self.lists)
        context.update(self.tables)

        jinja_env = jinja2.Environment(autoescape=True)

        # Заполняем шаблон
        self.doc.render(context, jinja_env)

        # Сохраняем результат
        self.doc.save(output_path)
        print(f'Document saved successfully. {output_path}')

        # for key in self.tables.keys():
        #    context[key] = [self.tables[key].columns.tolist()] + self.tables[key].values.tolist()
